"""简历 DOCX 生成器 — 根据原始简历 + AI 优化结果生成可下载的简历文件"""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def export_resume_docx(profile, optimization: dict, company: str, job_title: str) -> bytes:
    """生成优化后的简历 DOCX，返回字节流。

    profile: ResumeProfile (原始解析的简历)
    optimization: dict (包含 optimized_bullets, matched_skills, tailored_summary 等)
    company: 目标公司名（用于页眉）
    job_title: 目标岗位名（用于页眉）
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ===== 标题 =====
    _add_heading(doc, profile.name or "姓名", Pt(20))
    _add_paragraph(doc, f"求职意向：{job_title} @ {company}", Pt(11), bold=True)
    _add_paragraph(doc, "", Pt(6))

    # ===== 个人总结（使用 AI 优化版） =====
    summary = optimization.get("tailored_summary", "") or profile.summary or ""
    if summary:
        _add_section_title(doc, "个人总结")
        _add_paragraph(doc, summary, Pt(10.5))

    # ===== 技能 =====
    matched = optimization.get("matched_skills", []) or []
    all_skills = list(dict.fromkeys(matched + (profile.skills or [])))
    if all_skills:
        _add_section_title(doc, "专业技能")
        _add_paragraph(doc, " · ".join(all_skills), Pt(10.5))

    # ===== 工作经历（使用 AI 优化版 bullets） =====
    optimized_bullets = optimization.get("optimized_bullets", []) or []
    if optimized_bullets:
        _add_section_title(doc, "工作经历")
        for bullet in optimized_bullets:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet)
            run.font.size = Pt(10)
            run.font.name = "Microsoft YaHei"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    elif profile.work_experience:
        _add_section_title(doc, "工作经历")
        for exp in profile.work_experience:
            _add_paragraph(doc, f"{exp.title} | {exp.company} | {exp.duration}", Pt(10.5), bold=True)
            if exp.description:
                _add_paragraph(doc, exp.description, Pt(10))

    # ===== 项目经历 =====
    if profile.projects:
        _add_section_title(doc, "项目经历")
        for proj in profile.projects:
            techs = " | ".join(proj.technologies) if proj.technologies else ""
            _add_paragraph(doc, f"{proj.name}  {techs}", Pt(10.5), bold=True)
            if proj.description:
                _add_paragraph(doc, proj.description, Pt(10))

    # ===== 教育背景 =====
    if profile.education:
        _add_section_title(doc, "教育背景")
        for edu in profile.education:
            parts = [edu.institution, edu.degree, edu.major, edu.graduation]
            _add_paragraph(doc, " | ".join(p for p in parts if p), Pt(10.5))

    # ===== 保存 =====
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_heading(doc, text, size):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_paragraph(doc, text, size, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold:
        run.bold = True


def _add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    # 下划线分隔
    p_border = doc.add_paragraph()
    run_b = p_border.add_run("─" * 60)
    run_b.font.size = Pt(6)
    run_b.font.color.rgb = RGBColor(180, 180, 180)
