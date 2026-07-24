import io
import re
from app.models.resume import ResumeProfile, WorkExperience, Education, Project

SKILL_POOL = [
    "Python", "Java", "JavaScript", "TypeScript", "Go", "Rust", "C++", "C#",
    "React", "Vue", "Angular", "Next.js", "Node.js", "Django", "Flask",
    "FastAPI", "Spring Boot", "Express", "Redis", "MySQL", "PostgreSQL",
    "MongoDB", "Docker", "Kubernetes", "AWS", "Azure", "GCP", "Git",
    "Linux", "SQL", "GraphQL", "REST", "gRPC", "Kafka", "RabbitMQ",
    "HTML", "CSS", "Sass", "Tailwind", "TensorFlow", "PyTorch",
    "Spark", "Flink", "Hadoop", "Figma", "Webpack", "Vite",
]

TITLE_POOL = [
    "后端工程师", "前端工程师", "全栈工程师", "数据工程师",
    "算法工程师", "AI工程师", "产品经理", "项目经理",
    "架构师", "运维工程师", "测试工程师", "数据分析师",
    "DevOps", "SRE", "技术负责人",
]


PHONE_PATTERN = re.compile(
    r"(?:电话|手机|tel|phone|联系方式)[:：]?\s*"
    r"(\+?86[\s-]?)?(1[3-9]\d[\s-]?\d{4}[\s-]?\d{4})",
    re.IGNORECASE,
)
PHONE_BARE = re.compile(r"(?<!\d)(1[3-9]\d[\s-]?\d{4}[\s-]?\d{4})(?!\d)")
EMAIL_PATTERN = re.compile(
    r"(?:邮箱|email|e-mail|邮件)[:：]?\s*"
    r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
    re.IGNORECASE,
)
EMAIL_BARE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")

COMPANY_HINTS = re.compile(
    r"(?:公司[:：]|工作单位[:：]|employer[:：]|organization[:：])",
    re.IGNORECASE,
)
DURATION_PATTERN = re.compile(
    r"(\d{4}\.\d{1,2}|\d{4}-\d{1,2}|\d{4}/\d{1,2})\s*[-–—至到~]\s*"
    r"(\d{4}\.\d{1,2}|\d{4}-\d{1,2}|\d{4}/\d{1,2}|至今|现在|Present)",
)
SECTION_HEADERS = re.compile(
    r"^(?:##\s*)?"
    r"(技能|skill|技术栈|tech stack|工作经历|工作经[历验]|work experience|"
    r"项目经历|项目经[历验]|project|projects|教育经历|教育背[景]|education|"
    r"个人总结|自我评价|summary|求职意向|联系方式|个人信息)",
    re.IGNORECASE | re.MULTILINE,
)

# 中文字符占比阈值，低于此值认为文本提取可能损坏
_MIN_CHINESE_RATIO = 0.02


def _chinese_ratio(text: str) -> float:
    if not text:
        return 0
    chinese_chars = sum(1 for c in text if "\u4e00" <= c <= "\u9fff")
    return chinese_chars / len(text)


def _extract_text_from_bytes(data: bytes, filename: str = "") -> str:
    """根据文件类型提取文本内容。"""
    name_lower = filename.lower()
    content_type = "text"

    if name_lower.endswith(".pdf") or data[:4] == b"%PDF":
        content_type = "pdf"
    elif name_lower.endswith(".docx") or (len(data) > 4 and data[:2] == b"PK" and b"word/" in data[:2000]):
        content_type = "docx"
    elif name_lower.endswith(".doc"):
        content_type = "doc"

    if content_type == "pdf":
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            if not pages:
                raise ValueError("PDF 无可提取的文字，可能是扫描件或图片，建议上传文本格式简历")
            text = "\n".join(pages)
            if _chinese_ratio(text) < _MIN_CHINESE_RATIO:
                raise ValueError("PDF 文本提取后中文缺失，可能是字体编码问题。建议将简历另存为 .txt 或 .docx 后重新上传")
            return text
    elif content_type == "docx":
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        if not text.strip():
            raise ValueError("DOCX 无可提取的文字，请确认文件内容不为空")
        return text
    elif content_type == "doc":
        text = data.decode("utf-8", errors="ignore")
        text = re.sub(r"[^\u4e00-\u9fff\w\s\n\d\+\.\,\;\:\(\)\[\]\{\}\-]+", " ", text)
        text = re.sub(r"\s+", " ", text)
        if _chinese_ratio(text) < _MIN_CHINESE_RATIO:
            raise ValueError(".doc 旧格式不支持解析，请将简历另存为 .docx 或 .txt 后重新上传")
        return text
    else:
        return data.decode("utf-8", errors="ignore")


def _clean(text: str) -> str:
    return text.strip().rstrip("：:").strip()


def _extract_name(lines: list[str]) -> str:
    for line in lines[:5]:
        cleaned = _clean(line)
        if cleaned and not any(
            keyword in cleaned
            for keyword in (
                "简历", "resume", "技能", "skill", "经验",
                "工作", "教育", "项目", "电话", "邮箱", "求职",
            )
        ):
            if re.match(r"^[\u4e00-\u9fff·]{2,4}$", cleaned):
                return cleaned
            if re.match(r"^[A-Z][a-z]+\s+[A-Z][a-z]+$", cleaned):
                return cleaned
    return ""


def _extract_title(lines: list[str]) -> str:
    for line in lines[:8]:
        cleaned = _clean(line)
        for title in TITLE_POOL:
            if title in cleaned:
                return title
    return ""


def _extract_skills(text: str) -> list[str]:
    found = []
    for skill in SKILL_POOL:
        escaped = re.escape(skill)
        if re.search(rf"(?<![a-zA-Z]){escaped}(?![a-zA-Z])", text, re.IGNORECASE):
            found.append(skill)
    return found


def _extract_work_experience(lines: list[str]) -> list[WorkExperience]:
    experiences: list[WorkExperience] = []
    current: WorkExperience | None = None
    in_work = False
    desc_lines: list[str] = []

    for line in lines:
        if re.search(r"工作经历|工作经验|工作经[历验]|work\s*experience", line, re.IGNORECASE):
            in_work = True
            if current and current.company:
                current.description = "; ".join(desc_lines)
                experiences.append(current)
                desc_lines = []
            current = None
            continue
        if in_work:
            if re.search(r"项目经历|项目经[历验]|project|教育经历|教育背[景]|education|技能", line, re.IGNORECASE):
                if current and current.company:
                    current.description = "; ".join(desc_lines)
                    experiences.append(current)
                    desc_lines = []
                current = None
                in_work = False
                continue

            dur_match = DURATION_PATTERN.search(line)
            if dur_match and not COMPANY_HINTS.search(line):
                if current and current.company:
                    current.description = "; ".join(desc_lines)
                    experiences.append(current)
                    desc_lines = []

                current = WorkExperience(duration=dur_match.group(0))
                remainder = line.replace(dur_match.group(0), "").strip()
                parts = [p.strip() for p in remainder.split("|") if p.strip()]
                if len(parts) >= 2:
                    current.company = parts[0]
                    current.title = parts[1]
                elif remainder:
                    current.title = remainder
                continue

            if current:
                if COMPANY_HINTS.search(line):
                    current.company = line.split("：", 1)[-1].split(":", 1)[-1].strip()
                elif not current.title and not current.company and len(_clean(line)) > 1:
                    current.title = _clean(line)
                elif len(_clean(line)) > 1:
                    desc_lines.append(_clean(line))

    if current and current.company:
        current.description = "; ".join(desc_lines)
        experiences.append(current)

    return experiences


def _extract_projects(lines: list[str]) -> list[Project]:
    projects: list[Project] = []
    current: Project | None = None
    in_projects = False
    desc_lines: list[str] = []

    for line in lines:
        if re.search(r"项目经历|项目经[历验]|project", line, re.IGNORECASE):
            in_projects = True
            if current and current.name:
                current.description = "; ".join(desc_lines)
                projects.append(current)
                desc_lines = []
            current = None
            continue
        if in_projects:
            if re.search(r"教育经历|教育背[景]|education|技能|工作经历|工作经[历验]|work\s*experience", line, re.IGNORECASE):
                if current and current.name:
                    current.description = "; ".join(desc_lines)
                    projects.append(current)
                    desc_lines = []
                current = None
                in_projects = False
                continue

            cleaned = _clean(line)
            if not cleaned:
                continue

            if not current:
                current = Project(name=cleaned)
            else:
                desc_lines.append(cleaned)

    if current and current.name:
        current.description = "; ".join(desc_lines)
        projects.append(current)

    for proj in projects:
        proj.technologies = [s for s in SKILL_POOL if s.lower() in proj.description.lower()]

    return projects


def _extract_education(lines: list[str]) -> list[Education]:
    education: list[Education] = []
    current: Education | None = None
    in_edu = False

    for line in lines:
        if re.search(r"教育经历|教育背[景]|education", line, re.IGNORECASE):
            in_edu = True
            current = None
            continue
        if in_edu:
            if re.search(r"技能|skill|工作经历|工作经[历验]|work|项目经历|项目经[历验]|project", line, re.IGNORECASE):
                if current and current.institution:
                    education.append(current)
                current = None
                in_edu = False
                continue
            cleaned = _clean(line)
            if not cleaned:
                continue
            if not current:
                current = Education()
                parts = [p.strip() for p in cleaned.split("|") if p.strip()]
                if len(parts) >= 2:
                    current.institution = parts[0]
                    if len(parts) >= 3:
                        current.major = parts[1]
                        current.degree = parts[2]
                    else:
                        current.degree = parts[1]
                else:
                    current.institution = cleaned
            else:
                if "学位" in cleaned or "本科" in cleaned or "硕士" in cleaned or "博士" in cleaned:
                    current.degree = cleaned
                elif "专业" in cleaned:
                    current.major = cleaned.replace("专业", "").strip()

    if current and current.institution:
        education.append(current)

    return education



def _extract_phone(text: str) -> str:
    m = PHONE_PATTERN.search(text)
    if m:
        return re.sub(r"[\s-]", "", m.group(2))
    m = PHONE_BARE.search(text)
    return re.sub(r"[\s-]", "", m.group(1)) if m else ""


def _extract_email(text: str) -> str:
    m = EMAIL_PATTERN.search(text)
    if m:
        return m.group(1)
    m = EMAIL_BARE.search(text)
    return m.group(0) if m else ""


def _extract_gender(text: str) -> str:
    if re.search(r"性别[:：]?\s*男|男\s*[\|/]", text):
        return "男"
    if re.search(r"性别[:：]?\s*女|女\s*[\|/]", text):
        return "女"
    return ""


def _extract_birth(text: str) -> str:
    m = re.search(r"(?:出生|生日|birth)[:：]?\s*(\d{4}[\s./-]?\d{1,2}[\s./-]?\d{1,2})", text, re.I)
    return m.group(1) if m else ""


def _extract_location(text: str) -> str:
    m = re.search(r"(?:现居|所在地|城市|location)[:：]?\s*([一-鿿]{2,6}(?:市|省)?)", text, re.I)
    return m.group(1) if m else ""


def _extract_summary(lines: list[str]) -> str:
    for i, line in enumerate(lines[:20]):
        if re.search(r"个人总结|自我评价|summary|求职意向", line, re.IGNORECASE):
            if i + 1 < len(lines):
                return _clean(lines[i + 1])
    return ""


def parse_resume_text(text: str) -> ResumeProfile:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return ResumeProfile()

    if len(lines) < 3:
        joined = " ".join(lines)
        lines = [line.strip() for line in joined.replace("。", "。\n").splitlines() if line.strip()]
        if len(lines) < 3:
            return ResumeProfile()

    name = _extract_name(lines)
    title = _extract_title(lines)
    phone = _extract_phone(text)
    email = _extract_email(text)
    gender = _extract_gender(text)
    birth = _extract_birth(text)
    location = _extract_location(text)
    skills = _extract_skills(text)
    work_experience = _extract_work_experience(lines)
    projects = _extract_projects(lines)
    education = _extract_education(lines)
    summary = _extract_summary(lines)

    target_titles: list[str] = []
    for t in TITLE_POOL:
        if t in text:
            target_titles.append(t)

    profile = ResumeProfile(
        name=name,
        title=title,
        phone=phone,
        email=email,
        gender=gender,
        birth=birth,
        location=location,
        summary=summary,
        skills=skills,
        target_titles=target_titles,
        work_experience=work_experience,
        education=education,
        projects=projects,
    )

    # AI 优先解析：先用 AI 获取结构化数据，正则结果用于补充 AI 未覆盖的字段
    import logging
    logger = logging.getLogger("resume_parser")
    try:
        ai_profile = _ai_parse_resume(text, ["全部字段"])
        if ai_profile:
            # AI 结果优先，正则补充 AI 未提取到的字段
            profile = ResumeProfile(
                name=ai_profile.name or profile.name,
                title=ai_profile.title or profile.title,
                phone=getattr(ai_profile, 'phone', '') or getattr(profile, 'phone', ''),
                email=getattr(ai_profile, 'email', '') or getattr(profile, 'email', ''),
                gender=getattr(ai_profile, 'gender', '') or getattr(profile, 'gender', ''),
                birth=getattr(ai_profile, 'birth', '') or getattr(profile, 'birth', ''),
                location=getattr(ai_profile, 'location', '') or getattr(profile, 'location', ''),
                summary=ai_profile.summary or profile.summary,
                skills=ai_profile.skills if ai_profile.skills else profile.skills,
                target_titles=ai_profile.target_titles if ai_profile.target_titles else profile.target_titles,
                work_experience=ai_profile.work_experience if ai_profile.work_experience else profile.work_experience,
                education=ai_profile.education if ai_profile.education else profile.education,
                projects=ai_profile.projects if ai_profile.projects else profile.projects,
            )
            logger.info("AI 优先解析完成")
            return profile
    except Exception as e:
        logger.warning("AI 优先解析失败，使用正则结果: %s", e)

    # 正则兜底：AI 完全不可用时仍用正则
    missing = []
    if not profile.name: missing.append("姓名")
    if not profile.title: missing.append("职位")
    if not profile.work_experience: missing.append("工作经历")

    if missing:
        logger.info("正则解析缺失字段: %s，AI 已尝试但不可用", missing)

    return profile



def _ai_parse_resume(text: str, fields_missing: list[str]) -> ResumeProfile | None:
    """当正则解析缺失关键字段时，用 AI 从原文提取结构化简历数据。"""
    import logging
    logger = logging.getLogger("resume_parser")
    try:
        from app.services.ai_client import chat_json
    except Exception:
        return None

    system = """你是一位专业的简历解析器。从以下简历原文中提取结构化信息。

返回 JSON：
{
  "name": "姓名",
  "title": "当前职位/求职方向",
  "summary": "个人总结",
  "skills": ["技能1", "技能2"],
  "target_titles": ["期望岗位"],
  "work_experience": [
    {"company": "公司名", "title": "职位", "duration": "时间范围如2020.06-2023.08", "description": "工作描述"}
  ],
  "education": [
    {"institution": "学校", "degree": "学历", "major": "专业", "graduation": "毕业时间"}
  ],
  "projects": [
    {"name": "项目名", "description": "项目描述", "technologies": ["技术"]}
  ]
}

只返回 JSON，不要解释。如果某字段确实没有，用空字符串/空数组。"""

    # 只传前 5000 字符，避免 tokens 过大
    truncated = text[:5000]
    user = f"以下是一份简历的原文，请提取结构化信息。缺失字段：{', '.join(fields_missing)}。\n\n简历原文：\n{truncated}"

    try:
        data = chat_json(system, user)
    except Exception as e:
        logger.warning("AI 简历解析失败: %s", e)
        return None

    if not data or not isinstance(data, dict):
        return None

    # 构建 ResumeProfile
    experiences = []
    for exp in data.get("work_experience", []) or []:
        if isinstance(exp, dict):
            experiences.append(WorkExperience(
                company=str(exp.get("company", "")),
                title=str(exp.get("title", "")),
                duration=str(exp.get("duration", "")),
                description=str(exp.get("description", "")),
            ))

    education = []
    for edu in data.get("education", []) or []:
        if isinstance(edu, dict):
            education.append(Education(
                institution=str(edu.get("institution", "")),
                degree=str(edu.get("degree", "")),
                major=str(edu.get("major", "")),
                graduation=str(edu.get("graduation", "")),
            ))

    projects = []
    for proj in data.get("projects", []) or []:
        if isinstance(proj, dict):
            projects.append(Project(
                name=str(proj.get("name", "")),
                description=str(proj.get("description", "")),
                technologies=[str(t) for t in (proj.get("technologies", []) or [])],
            ))

    return ResumeProfile(
        name=str(data.get("name", "")),
        title=str(data.get("title", "")),
        summary=str(data.get("summary", "")),
        skills=[str(s) for s in (data.get("skills", []) or [])],
        target_titles=[str(s) for s in (data.get("target_titles", []) or [])],
        work_experience=experiences,
        education=education,
        projects=projects,
    )


def parse_resume_bytes(data: bytes, filename: str = "") -> ResumeProfile:
    text = _extract_text_from_bytes(data, filename)
    return parse_resume_text(text)
